"""
인프라보안 전용 템플릿 시드 스크립트
- sample/ 디렉토리의 문서를 기반으로 템플릿 생성
- {{변수}} 치환을 통한 자동화 지원
"""

import io
import os
import re
import tempfile
from pathlib import Path

import minio_client
from models import DocumentTemplate, TemplateBundle, TemplateBundleItem, User
from template_package_utils import (
    extract_placeholders_from_office_package,
    replace_text_in_office_package,
)

SAMPLE_DIR = Path(__file__).resolve().parent.parent.parent / "sample"
BUNDLE_NAME = "인프라보안 전용 템플릿"

# 발주서에서 추출한 실제 값 → 변수 매핑 (치환 대상)
REPLACEMENTS = {
    "PO-20260319-0043": "{{발주번호}}",
    "[베트남 하노이 센터] Extreme 7520 스위치 1대 구매 건": "{{발주명}}",
    "2026/03/19": "{{발주일자}}",
    "2026/03/30": "{{입고일자}}",
    "2026/03/31": "{{검수일자}}",
    "2026-03-19": "{{발주일자}}",
    "2026-03-30": "{{입고일자}}",
    "2026-03-31": "{{검수일자}}",
    "SM022609Q-40056": "{{시리얼번호}}",
    "아이클라우드 주식회사": "{{공급사}}",
    "아이클라우드": "{{공급사_약칭}}",
    "가산2IDC": "{{납품장소}}",
    "전진호": "{{공급사담당자}}",
    "7520-48Y-8C-AC-F": "{{모델명}}",
    "2027-03-31": "{{유지보수종료일}}",
}

DOCX_TEXT_REPLACEMENTS = [
    ("[베트남 하노이 센터] Extreme 7520 스위치 1대 구매 건", "{{발주명}}"),
    ("PO-20260319-0043", "{{발주번호}}"),
    ("2026/ 03 /19", "{{발주일자_슬래시공백}}"),
    ("2026/ 03 /30", "{{입고일자_슬래시공백}}"),
    ("2026/ 03 /31", "{{검수일자_슬래시공백}}"),
    ("2026/ 03/31", "{{검수일자_슬래시공백}}"),
    ("2026/03/19", "{{발주일자}}"),
    ("2026/03/30", "{{입고일자}}"),
    ("2026/03/31", "{{검수일자}}"),
    ("2026년 03월 19일", "{{발주일자_한글}}"),
    ("2026년 03월 30일", "{{입고일자_한글}}"),
    ("2026년 03월 31일", "{{검수일자_한글}}"),
    ("2026 년 03 월 31 일", "{{검수일자_한글띄어쓰기}}"),
    ("03/31/2026", "{{검수일자_US}}"),
    ("가산2\nIDC", "{{납품장소}}"),
    ("가산2 IDC", "{{납품장소}}"),
    ("가산2", "{{납품장소}}"),
    ("아이클라우드 주식회사", "{{공급사}}"),
    ("아이클라우드", "{{공급사_약칭}}"),
    ("7520-48Y-8C-AC-F 1식", "{{모델명}} {{수량}}식"),
    ("1ea", "{{수량}}ea"),
    ("SM022609Q-40056", "{{시리얼번호}}"),
    ("7520-48Y-8C-AC-F", "{{모델명}}"),
]

BUNDLE_VARIABLES = [
    {"key": "발주번호", "label": "발주번호 (PO#)", "type": "text", "section": "발주정보"},
    {"key": "발주명", "label": "발주명", "type": "text", "section": "발주정보"},
    {"key": "발주일자", "label": "발주일자", "type": "text", "section": "발주정보"},
    {"key": "납품기한", "label": "납품기한", "type": "text", "section": "발주정보"},
    {"key": "제조사", "label": "제조사", "type": "text", "section": "장비정보"},
    {"key": "모델명", "label": "모델명", "type": "text", "section": "장비정보"},
    {"key": "수량", "label": "수량 (EA)", "type": "text", "section": "장비정보"},
    {"key": "시리얼번호", "label": "시리얼번호 (S/N)", "type": "text", "section": "장비정보"},
    {"key": "OS버전", "label": "OS 버전", "type": "text", "section": "장비정보"},
    {"key": "입고일자", "label": "입고일자", "type": "text", "section": "납품검수"},
    {"key": "검수일자", "label": "검수일자", "type": "text", "section": "납품검수"},
    {"key": "납품장소", "label": "납품장소 (IDC)", "type": "text", "section": "납품검수"},
    {"key": "공급사", "label": "공급사 (정식명칭)", "type": "text", "section": "납품검수"},
    {"key": "공급사_약칭", "label": "공급사 (약칭)", "type": "text", "section": "납품검수"},
    {"key": "공급사담당자", "label": "공급사 담당자", "type": "text", "section": "납품검수"},
    {"key": "유지보수수량", "label": "유지보수 수량", "type": "text", "section": "유지보수"},
    {"key": "유지보수종료일", "label": "유지보수 종료일", "type": "text", "section": "유지보수"},
    {"key": "출입자1_회사명", "label": "출입자1 회사명", "type": "text", "section": "IDC출입"},
    {"key": "출입자1_이름", "label": "출입자1 이름", "type": "text", "section": "IDC출입"},
    {"key": "출입자1_직책", "label": "출입자1 직책", "type": "text", "section": "IDC출입"},
    {"key": "출입자1_연락처", "label": "출입자1 연락처", "type": "text", "section": "IDC출입"},
    {"key": "출입자2_회사명", "label": "출입자2 회사명", "type": "text", "section": "IDC출입"},
    {"key": "출입자2_이름", "label": "출입자2 이름", "type": "text", "section": "IDC출입"},
    {"key": "출입자2_직책", "label": "출입자2 직책", "type": "text", "section": "IDC출입"},
    {"key": "출입자2_연락처", "label": "출입자2 연락처", "type": "text", "section": "IDC출입"},
]

# 각 템플릿 문서 정의
TEMPLATE_DEFS = [
    {
        "sample_file": "7520-48Y-8C-AC-F_검수확인서_PO-20260319-0043_1EA.docx",
        "name": "검수확인서",
        "display_name": "검수확인서",
        "output_pattern": "{{모델명}}_검수확인서_{{발주번호}}_{{수량}}EA",
        "description": "장비 입고 후 검수 결과를 기록하는 확인서",
    },
    {
        "sample_file": "7520-48Y-8C-AC-F_현장검수확인서_PO-20260319-0043_1EA.docx",
        "name": "현장검수확인서",
        "display_name": "현장검수확인서",
        "output_pattern": "{{모델명}}_현장검수확인서_{{발주번호}}_{{수량}}EA",
        "description": "현장 검수 시험 절차 및 결과 리포트",
    },
    {
        "sample_file": "20260331_보안장비_무결성체크_점검.docx",
        "name": "보안장비 무결성체크 점검",
        "display_name": "보안장비 무결성체크",
        "output_pattern": "20260331_보안장비_무결성체크_점검",
        "description": "어플라이언스 장비 무결성 점검 체크리스트",
    },
    {
        "sample_file": "NBP 입고일정공유파일.xlsx",
        "name": "NBP 입고일정공유파일",
        "display_name": "NBP 입고일정",
        "output_pattern": "NBP 입고일정공유파일",
        "description": "NBP 입고 일정 공유 파일",
    },
    {
        "sample_file": "CMDB_PO-20260319-0043.xlsm",
        "name": "CMDB",
        "display_name": "CMDB",
        "output_pattern": "CMDB_{{발주번호}}",
        "description": "네트워크 H/W 정보 입력 양식 (CMDB)",
    },
]

# XLS -> XLSX 변환이 필요한 파일 (openpyxl로 새로 생성)
XLSX_FROM_SCRATCH = [
    {
        "name": "납품확인서",
        "display_name": "납품확인서",
        "sample_file": "납품확인서_Extreme 7520 스위치 1대 구매_260330.xls",
        "output_pattern": "납품확인서",
        "description": "납품/검수 확인서",
    },
    {
        "name": "IDC 출입명단",
        "display_name": "IDC 출입명단",
        "sample_file": "IDC 출입명단.xls",
        "output_pattern": "IDC 출입명단",
        "description": "IDC 출입 및 권한 신청서",
    },
]


def _replace_in_paragraph(paragraph, replacements: dict) -> bool:
    """단락 내 텍스트를 치환. 서식은 첫 번째 run의 것을 유지."""
    full_text = "".join(run.text for run in paragraph.runs)
    changed = False
    for old, new in replacements.items():
        if old in full_text:
            full_text = full_text.replace(old, new)
            changed = True
    if changed and paragraph.runs:
        paragraph.runs[0].text = full_text
        for run in paragraph.runs[1:]:
            run.text = ""
    return changed


def _iter_table_paragraphs(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                yield paragraph
            for nested_table in cell.tables:
                yield from _iter_table_paragraphs(nested_table)


def _iter_doc_paragraphs(doc):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        yield from _iter_table_paragraphs(table)
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            yield paragraph
        for table in section.header.tables:
            yield from _iter_table_paragraphs(table)
        for paragraph in section.footer.paragraphs:
            yield paragraph
        for table in section.footer.tables:
            yield from _iter_table_paragraphs(table)


def _create_docx_template(sample_path: str, replacements: dict) -> bytes:
    """DOCX 샘플 파일을 열어 특정 값을 {{변수}}로 치환한 뒤 bytes 반환."""
    from docx import Document

    doc = Document(sample_path)
    ordered_replacements = dict(DOCX_TEXT_REPLACEMENTS)

    for paragraph in _iter_doc_paragraphs(doc):
        _replace_in_paragraph(paragraph, ordered_replacements)

    output = io.BytesIO()
    doc.save(output)

    package_replacements = dict(DOCX_TEXT_REPLACEMENTS)
    package_replacements.update(replacements)
    return replace_text_in_office_package(
        output.getvalue(),
        package_replacements,
        xml_prefixes=("word/",),
        escape_xml=False,
    )


def _create_xlsx_template_from_sample(sample_path: str, replacements: dict) -> bytes:
    """XLSX/XLSM 샘플 파일을 열어 셀 값을 치환한 뒤 XLSX bytes 반환."""
    import openpyxl

    wb = openpyxl.load_workbook(sample_path)
    for sheet in wb.worksheets:
        for row in sheet.iter_rows():
            for cell in row:
                if cell.value and isinstance(cell.value, str):
                    val = cell.value
                    for old, new in replacements.items():
                        val = val.replace(old, new)
                    cell.value = val
    buf = io.BytesIO()
    wb.save(buf)
    return replace_text_in_office_package(
        buf.getvalue(),
        replacements,
        xml_prefixes=("xl/",),
        escape_xml=False,
    )


def _convert_xls_sample_to_clean_xlsx(sample_path: str) -> bytes:
    import jdk4py
    import openpyxl

    os.environ.setdefault("JAVA_HOME", str(jdk4py.JAVA_HOME))

    import asposecells

    if not asposecells.isJVMStarted():
        asposecells.startJVM()

    pkg = asposecells.JPackage("com").aspose.cells
    workbook = pkg.Workbook(sample_path)

    with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as tmp_output:
        output_path = tmp_output.name

    try:
        workbook.save(output_path, pkg.SaveFormat.XLSX)
        converted = Path(output_path).read_bytes()
    finally:
        Path(output_path).unlink(missing_ok=True)

    wb = openpyxl.load_workbook(io.BytesIO(converted))
    if "Evaluation Warning" in wb.sheetnames:
        wb.remove(wb["Evaluation Warning"])

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _create_nbp_schedule_template(sample_path: str) -> bytes:
    import openpyxl

    wb = openpyxl.load_workbook(sample_path)
    ws = wb[wb.sheetnames[0]]
    ws["A2"] = "{{제조사}}"
    ws["B2"] = "{{공급사_약칭}}"
    ws["C2"] = "{{발주번호}}"
    ws["D2"] = "{{모델명}}"
    ws["E2"] = "{{수량}}EA"
    ws["F2"] = "{{발주일자}}"
    ws["G2"] = "{{입고일자}}"
    ws["H2"] = "{{입고일자}}"
    ws["I2"] = "{{납품장소}}"

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _create_delivery_confirmation_template(sample_path: str) -> bytes:
    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(_convert_xls_sample_to_clean_xlsx(sample_path)))
    ws = wb[wb.sheetnames[0]]
    ws["C7"] = "{{입고일자_대시}}"
    ws["C11"] = "{{발주명}}"
    ws["C15"] = "{{모델명}}"
    ws["D15"] = "{{시리얼번호}}"
    ws["E15"] = "{{수량}}"
    ws["E16"] = "{{수량}}"
    ws["E17"] = "{{수량}}"
    ws["E18"] = "{{유지보수수량}}"
    ws["E32"] = "{{입고일자_한글여백}}"

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _create_idc_access_xls_template(sample_path: str) -> bytes:
    import xlrd
    from xlutils.copy import copy as xl_copy

    book = xlrd.open_workbook(sample_path, formatting_info=True)
    writable = xl_copy(book)
    sheet = writable.get_sheet(0)
    sheet.set_panes_frozen(True)
    sheet.set_horz_split_pos(8)
    sheet.set_horz_split_first_visible(8)
    sheet.set_remove_splits(True)

    buf = io.BytesIO()
    writable.save(buf)
    return buf.getvalue()


def _xls_colour_to_rgb(book, colour_index: int) -> str | None:
    if colour_index in (None, 0x7FFF):
        return None
    rgb = book.colour_map.get(colour_index)
    if not rgb:
        return None
    return "".join(f"{part:02X}" for part in rgb)


def _xls_border_style(style_code: int) -> str | None:
    return {
        0: None,
        1: "thin",
        2: "medium",
        3: "dashed",
        4: "dotted",
        5: "thick",
        6: "double",
        7: "hair",
        8: "mediumDashed",
        9: "dashDot",
        10: "mediumDashDot",
        11: "dashDotDot",
        12: "mediumDashDotDot",
        13: "slantDashDot",
    }.get(style_code)


def _openpyxl_side(book, line_style: int, colour_index: int):
    from openpyxl.styles import Color, Side

    style = _xls_border_style(line_style)
    if not style:
        return Side()
    color = _xls_colour_to_rgb(book, colour_index)
    return Side(style=style, color=Color(rgb=color) if color else None)


def _apply_xls_cell_style(book, xf, cell):
    from openpyxl.styles import Alignment, Border, Color, Font, PatternFill

    font = book.font_list[xf.font_index]
    font_color = _xls_colour_to_rgb(book, font.colour_index)
    cell.font = Font(
        name=font.name,
        size=(font.height / 20) if font.height else None,
        bold=bool(font.bold),
        italic=bool(font.italic),
        underline="single" if font.underline_type else None,
        strike=bool(font.struck_out),
        color=Color(rgb=font_color) if font_color else None,
    )

    background = xf.background
    fill_color = _xls_colour_to_rgb(book, background.pattern_colour_index)
    if background.fill_pattern and fill_color:
        cell.fill = PatternFill(fill_type="solid", fgColor=fill_color, bgColor=fill_color)

    border = xf.border
    cell.border = Border(
        left=_openpyxl_side(book, border.left_line_style, border.left_colour_index),
        right=_openpyxl_side(book, border.right_line_style, border.right_colour_index),
        top=_openpyxl_side(book, border.top_line_style, border.top_colour_index),
        bottom=_openpyxl_side(book, border.bottom_line_style, border.bottom_colour_index),
    )

    alignment = xf.alignment
    cell.alignment = Alignment(
        horizontal={
            1: "left",
            2: "center",
            3: "right",
            5: "justify",
            6: "centerContinuous",
            7: "distributed",
        }.get(alignment.hor_align),
        vertical={
            0: "top",
            1: "center",
            2: "bottom",
            3: "justify",
            4: "distributed",
        }.get(alignment.vert_align),
        wrap_text=bool(alignment.text_wrapped),
        shrink_to_fit=bool(alignment.shrink_to_fit),
        text_rotation=alignment.rotation if 0 <= alignment.rotation <= 180 else 0,
    )

    fmt = book.format_map.get(xf.format_key)
    if fmt and fmt.format_str:
        cell.number_format = fmt.format_str


def _convert_xls_sample_to_xlsx(sample_path: str, replacements: dict) -> bytes:
    import openpyxl
    import xlrd

    book = xlrd.open_workbook(sample_path, formatting_info=True)
    out_wb = openpyxl.Workbook()
    out_wb.remove(out_wb.active)

    for sheet in book.sheets():
        ws = out_wb.create_sheet(title=sheet.name[:31] or "Sheet1")

        for col_idx, col_info in sheet.colinfo_map.items():
            ws.column_dimensions[openpyxl.utils.get_column_letter(col_idx + 1)].width = col_info.width / 256

        for row_idx, row_info in sheet.rowinfo_map.items():
            ws.row_dimensions[row_idx + 1].height = row_info.height / 20
            ws.row_dimensions[row_idx + 1].hidden = bool(row_info.hidden)

        for row_idx in range(sheet.nrows):
            for col_idx in range(sheet.ncols):
                source = sheet.cell(row_idx, col_idx)
                target = ws.cell(row_idx + 1, col_idx + 1)
                value = source.value
                if isinstance(value, str):
                    for old, new in replacements.items():
                        value = value.replace(old, new)
                if source.ctype != xlrd.XL_CELL_EMPTY:
                    target.value = value

                xf = book.xf_list[sheet.cell_xf_index(row_idx, col_idx)]
                _apply_xls_cell_style(book, xf, target)

        for rlo, rhi, clo, chi in sheet.merged_cells:
            ws.merge_cells(
                start_row=rlo + 1,
                end_row=rhi,
                start_column=clo + 1,
                end_column=chi,
            )

    buf = io.BytesIO()
    out_wb.save(buf)
    return buf.getvalue()


def _set_sheet_cell(ws, row: int, col: int, value: str) -> None:
    ws.cell(row=row, column=col).value = value


def _apply_idc_access_placeholders(xlsx_data: bytes) -> bytes:
    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(xlsx_data))
    ws = wb.active
    _set_sheet_cell(ws, 9, 1, "")
    _set_sheet_cell(ws, 9, 2, "{{출입자1_회사명}}")
    _set_sheet_cell(ws, 9, 3, "{{출입자1_이름}}")
    _set_sheet_cell(ws, 9, 4, "{{출입자1_직책}}")
    _set_sheet_cell(ws, 9, 5, "{{출입자1_연락처}}")
    _set_sheet_cell(ws, 10, 1, "")
    _set_sheet_cell(ws, 10, 2, "{{출입자2_회사명}}")
    _set_sheet_cell(ws, 10, 3, "{{출입자2_이름}}")
    _set_sheet_cell(ws, 10, 4, "{{출입자2_직책}}")
    _set_sheet_cell(ws, 10, 5, "{{출입자2_연락처}}")
    for col in range(1, 6):
        _set_sheet_cell(ws, 11, col, "")

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _upsert_template(
    db_session,
    admin: User,
    template_name: str,
    description: str,
    file_type: str,
    template_filename: str,
    data: bytes,
):
    object_key = f"documents/templates/{admin.id}/bundles/{template_filename}"
    content_type = {
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "xlsm": "application/vnd.ms-excel.sheet.macroEnabled.12",
        "xls": "application/vnd.ms-excel",
    }.get(file_type, "application/octet-stream")

    if not minio_client.upload_file(object_key, data, content_type):
        print(f"[Seed] MinIO 업로드 실패: {template_filename}")
        return None

    variables = _extract_variables(data, file_type)
    tpl = (
        db_session.query(DocumentTemplate)
        .filter(DocumentTemplate.name == template_name)
        .first()
    )
    if not tpl:
        tpl = DocumentTemplate(name=template_name, created_by=admin.id)
        db_session.add(tpl)

    tpl.description = description or None
    tpl.file_type = file_type
    tpl.original_filename = template_filename
    tpl.minio_object_key = object_key
    tpl.file_size = len(data)
    tpl.variables = variables

    db_session.flush()
    return tpl


def seed_infra_security_bundle(db_session):
    """인프라보안 전용 템플릿 번들 메타데이터만 동기화."""

    admin = db_session.query(User).filter(User.is_admin.is_(True)).first()
    if not admin:
        print("[Seed] 관리자 계정을 찾을 수 없습니다. 번들 시드 스킵.")
        return

    bundle = (
        db_session.query(TemplateBundle)
        .filter(TemplateBundle.name == BUNDLE_NAME)
        .first()
    )
    is_new_bundle = bundle is None
    if not bundle:
        bundle = TemplateBundle(name=BUNDLE_NAME, created_by=admin.id)
        db_session.add(bundle)
        db_session.flush()

    bundle.template_folder = BUNDLE_NAME
    bundle.description = "발주서 기반으로 인프라보안 장비 관련 문서를 일괄 생성하는 전용 템플릿입니다."
    bundle.variables = BUNDLE_VARIABLES

    print(f"[Seed] '{BUNDLE_NAME}' 번들 {'생성' if is_new_bundle else '동기화'} 중...")

    existing_items = {item.display_name: item for item in bundle.items}
    template_defs = TEMPLATE_DEFS + XLSX_FROM_SCRATCH
    desired_display_names = {tdef["display_name"] for tdef in template_defs}

    for display_name, item in list(existing_items.items()):
        if display_name not in desired_display_names:
            db_session.delete(item)

    linked_count = 0
    for order, tdef in enumerate(template_defs):
        tpl = (
            db_session.query(DocumentTemplate)
            .filter(DocumentTemplate.name == f"[인프라보안] {tdef['name']}")
            .first()
        )
        if not tpl:
            print(f"[Seed]   템플릿 없음, 연결 스킵: {tdef['name']}")
            continue

        if tpl.folder_name != BUNDLE_NAME:
            tpl.folder_name = BUNDLE_NAME

        item = existing_items.get(tdef["display_name"])
        if not item:
            item = TemplateBundleItem(bundle_id=bundle.id, display_name=tdef["display_name"])
            db_session.add(item)

        item.template_id = tpl.id
        item.output_name_pattern = tdef.get("output_pattern", tdef["display_name"])
        item.order = order
        linked_count += 1
        print(f"[Seed]   기존 템플릿 연결: {tdef['name']} ({tpl.file_type})")

    db_session.commit()
    print(f"[Seed] '{BUNDLE_NAME}' 번들 동기화 완료 ({linked_count}개 템플릿 연결)")


def _extract_variables(data: bytes, file_type: str) -> list:
    """파일에서 {{변수}} 패턴을 추출."""
    if file_type == "docx":
        keys = extract_placeholders_from_office_package(data, ("word/",))
    elif file_type in {"xlsx", "xlsm"}:
        keys = extract_placeholders_from_office_package(data, ("xl/",))
    elif file_type == "xls":
        import xlrd

        pattern = re.compile(r"\{\{([^{}]+)\}\}")
        keys = set()
        book = xlrd.open_workbook(file_contents=data, formatting_info=True)
        for sheet in book.sheets():
            for row_idx in range(sheet.nrows):
                for col_idx in range(sheet.ncols):
                    value = sheet.cell_value(row_idx, col_idx)
                    if not isinstance(value, str):
                        continue
                    for match in pattern.finditer(value):
                        keys.add(match.group(1).strip())
    else:
        keys = set()
    return [{"key": key, "label": key, "type": "text"} for key in sorted(keys)]
