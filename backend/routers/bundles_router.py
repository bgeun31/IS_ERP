import io
import json
import re
import zipfile
from copy import copy, deepcopy
from typing import List
from urllib.parse import quote

from fastapi import APIRouter, Depends, File, HTTPException, Request, UploadFile
from fastapi.responses import Response
from sqlalchemy.orm import Session

import minio_client
from auth import get_current_user
from database import get_db
from models import DocumentRecord, DocumentTemplate, TemplateBundle, TemplateBundleItem, User
from purchase_order_parser import parse_purchase_order_pdf
from schemas import (
    BundlePurchaseOrderExtractResponse,
    TemplateBundleItemResponse,
    TemplateBundleResponse,
)
from template_package_utils import replace_text_in_office_package

router = APIRouter(prefix="/api/documents/bundles", tags=["bundles"])

DOCX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
XLSX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
)
XLSM_CONTENT_TYPE = "application/vnd.ms-excel.sheet.macroEnabled.12"
XLS_CONTENT_TYPE = "application/vnd.ms-excel"


def _content_type(file_type: str) -> str:
    return {
        "docx": DOCX_CONTENT_TYPE,
        "xlsx": XLSX_CONTENT_TYPE,
        "xlsm": XLSM_CONTENT_TYPE,
        "xls": XLS_CONTENT_TYPE,
    }.get(file_type, "application/octet-stream")


def _sanitize_filename(value: str) -> str:
    return re.sub(r'[\\/*?:"<>|]', "_", value)


def _compact_purchase_title(field_values: dict) -> str:
    title = str(field_values.get("발주명", "") or "").strip()
    if not title:
        return ""
    title = re.sub(r"^\[[^\]]+\]\s*", "", title)
    title = re.sub(r"\s*건$", "", title).strip()
    return title


def _compact_date_yyMMdd(value: str | None) -> str:
    raw = str(value or "").strip()
    match = re.search(r"(20\d{2})[/-](\d{2})[/-](\d{2})", raw)
    if match:
        return f"{match.group(1)[2:]}{match.group(2)}{match.group(3)}"
    return raw


def _compact_date_yyyyMMdd(value: str | None) -> str:
    raw = str(value or "").strip()
    match = re.search(r"(20\d{2})[/-](\d{2})[/-](\d{2})", raw)
    if match:
        return f"{match.group(1)}{match.group(2)}{match.group(3)}"
    return raw


def _split_date_parts(value: str | None) -> tuple[str, str, str] | None:
    raw = str(value or "").strip()
    match = re.search(r"(20\d{2})[/-](\d{1,2})[/-](\d{1,2})", raw)
    if not match:
        return None
    return match.group(1), f"{int(match.group(2)):02d}", f"{int(match.group(3)):02d}"


def _resolve_output_name(item: TemplateBundleItem, field_values: dict) -> str:
    quantity = str(field_values.get("수량", "") or "").strip()

    if item.display_name == "검수확인서":
        if quantity:
            return f"{field_values.get('모델명', '')}_검수확인서_{field_values.get('발주번호', '')}_{quantity}EA"
        return f"{field_values.get('모델명', '')}_검수확인서_{field_values.get('발주번호', '')}".strip("_")

    if item.display_name == "현장검수확인서":
        if quantity:
            return f"{field_values.get('모델명', '')}_현장검수확인서_{field_values.get('발주번호', '')}_{quantity}EA"
        return f"{field_values.get('모델명', '')}_현장검수확인서_{field_values.get('발주번호', '')}".strip("_")

    if item.display_name == "IDC 출입명단":
        return "IDC 출입명단"

    if item.display_name == "NBP 입고일정":
        return "NBP 입고일정공유파일"

    if item.display_name == "납품확인서":
        compact_title = _compact_purchase_title(field_values) or str(field_values.get("발주명", "") or "").strip()
        compact_date = _compact_date_yyMMdd(
            field_values.get("입고일자") or field_values.get("검수일자") or field_values.get("발주일자")
        )
        if compact_title and compact_date:
            return f"납품확인서_{compact_title}_{compact_date}"
        if compact_title:
            return f"납품확인서_{compact_title}"
        return "납품확인서"

    if item.display_name == "보안장비 무결성체크":
        compact_date = _compact_date_yyyyMMdd(field_values.get("검수일자") or field_values.get("입고일자"))
        if compact_date:
            return f"{compact_date}_보안장비_무결성체크_점검"
        return "보안장비_무결성체크_점검"

    output_name = item.output_name_pattern or item.display_name
    for key, value in field_values.items():
        output_name = output_name.replace(f"{{{{{key}}}}}", str(value))
    return output_name


def _bundle_to_response(b: TemplateBundle) -> TemplateBundleResponse:
    items = _ordered_bundle_items(b)
    return TemplateBundleResponse(
        id=b.id,
        name=b.name,
        template_folder=b.template_folder,
        description=b.description,
        variables=b.variables or [],
        items=[
            TemplateBundleItemResponse(
                id=item.id,
                template_id=item.template_id,
                display_name=item.display_name,
                output_name_pattern=item.output_name_pattern,
                file_type=item.template.file_type if item.template else None,
                order=item.order,
            )
            for item in items
        ],
        created_at=b.created_at,
        created_by_username=b.creator.username if b.creator else None,
    )


def _display_name_from_template(template: DocumentTemplate) -> str:
    name = (template.name or "").strip()
    match = re.match(r"^\[[^\]]+\]\s*(.+)$", name)
    return match.group(1).strip() if match else name


def _bundle_item_defaults(bundle: TemplateBundle) -> dict[str, tuple[str, int]]:
    if bundle.name != "인프라보안 전용 템플릿":
        return {}

    from seeds.infra_security import TEMPLATE_DEFS, XLSX_FROM_SCRATCH

    defaults: dict[str, tuple[str, int]] = {}
    for index, definition in enumerate(TEMPLATE_DEFS + XLSX_FROM_SCRATCH):
        defaults[definition["display_name"]] = (
            definition.get("output_pattern", definition["display_name"]),
            index,
        )
    return defaults


def _ordered_bundle_items(bundle: TemplateBundle) -> list[TemplateBundleItem]:
    return sorted(bundle.items, key=lambda item: (item.order, item.display_name, item.id))


def _sync_bundle_items_from_folder(db: Session, bundle: TemplateBundle) -> TemplateBundle:
    folder_name = (bundle.template_folder or "").strip()
    if not folder_name:
        return bundle

    templates = (
        db.query(DocumentTemplate)
        .filter(DocumentTemplate.folder_name == folder_name)
        .order_by(DocumentTemplate.name.asc(), DocumentTemplate.created_at.asc())
        .all()
    )
    item_defaults = _bundle_item_defaults(bundle)
    existing_items = {item.display_name: item for item in bundle.items}
    seen_display_names: set[str] = set()
    changed = False

    def sort_key(template: DocumentTemplate) -> tuple[int, str]:
        display_name = _display_name_from_template(template)
        return item_defaults.get(display_name, ("", 9999))[1], display_name

    for template in sorted(templates, key=sort_key):
        display_name = _display_name_from_template(template)
        if not display_name or display_name in seen_display_names:
            continue
        seen_display_names.add(display_name)

        output_pattern, default_order = item_defaults.get(display_name, (display_name, len(seen_display_names) - 1))
        item = existing_items.get(display_name)
        if not item:
            item = TemplateBundleItem(bundle_id=bundle.id, display_name=display_name)
            db.add(item)
            changed = True

        if item.template_id != template.id:
            item.template_id = template.id
            changed = True
        if item.output_name_pattern != output_pattern:
            item.output_name_pattern = output_pattern
            changed = True
        if item.order != default_order:
            item.order = default_order
            changed = True

    for display_name, item in list(existing_items.items()):
        if display_name in seen_display_names:
            continue
        db.delete(item)
        changed = True

    if changed:
        db.commit()
        db.refresh(bundle)

    return bundle


@router.get("", response_model=List[TemplateBundleResponse])
def list_bundles(
    db: Session = Depends(get_db),
    _: User = Depends(get_current_user),
):
    bundles = db.query(TemplateBundle).order_by(TemplateBundle.created_at.desc()).all()
    return [_bundle_to_response(_sync_bundle_items_from_folder(db, b)) for b in bundles]


@router.get("/{bundle_id}", response_model=TemplateBundleResponse)
def get_bundle(
    bundle_id: int,
    db: Session = Depends(get_db),
    _: User = Depends(get_current_user),
):
    b = db.query(TemplateBundle).filter(TemplateBundle.id == bundle_id).first()
    if not b:
        raise HTTPException(status_code=404, detail="번들을 찾을 수 없습니다")
    return _bundle_to_response(_sync_bundle_items_from_folder(db, b))


def _render_docx(template_data: bytes, field_values: dict) -> bytes:
    from docx import Document

    doc = Document(io.BytesIO(template_data))
    _replace_docx_paragraph_text(doc, _build_docx_legacy_date_replacements(field_values))

    output = io.BytesIO()
    doc.save(output)

    replacements = _build_scalar_replacements(field_values, serial_separator=", ")
    replacements.update(_build_docx_legacy_date_replacements(field_values))

    return replace_text_in_office_package(
        output.getvalue(),
        replacements,
        xml_prefixes=("word/",),
        escape_xml=True,
    )


def _render_xlsx(
    template_data: bytes,
    field_values: dict,
    *,
    expand_serial_rows: bool = True,
    keep_vba: bool = False,
) -> bytes:
    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(template_data), keep_vba=keep_vba)
    serial_numbers = _extract_serial_numbers(field_values)

    if expand_serial_rows and len(serial_numbers) > 1:
        for ws in wb.worksheets:
            _expand_xlsx_serial_rows(ws, serial_numbers)

    scalar_replacements = _build_scalar_replacements(field_values, serial_separator="\n")
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if not isinstance(cell.value, str):
                    continue
                value = cell.value
                for placeholder, replacement in scalar_replacements.items():
                    value = value.replace(placeholder, "" if replacement is None else str(replacement))
                cell.value = value

    output = io.BytesIO()
    wb.save(output)
    return output.getvalue()


def _parse_excel_date(value: str | None):
    from datetime import datetime

    raw = str(value or "").strip()
    if not raw:
        return None
    for fmt in ("%Y/%m/%d", "%Y-%m-%d"):
        try:
            return datetime.strptime(raw, fmt)
        except ValueError:
            continue
    return None


def _render_cmdb_xlsm(template_data: bytes, field_values: dict) -> bytes:
    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(template_data), keep_vba=True)
    ws = wb["NETWORK"] if "NETWORK" in wb.sheetnames else wb.active

    serial_numbers = [serial for serial in _extract_serial_numbers(field_values) if serial]
    quantity = str(field_values.get("수량", "") or "").strip()
    quantity_int = int(quantity) if quantity.isdigit() else 1
    row_count = max(len(serial_numbers), quantity_int, 1)
    base_row = 14

    if row_count > 1:
        ws.insert_rows(base_row + 1, row_count - 1)
        for offset in range(1, row_count):
            _copy_xlsx_row(ws, base_row, base_row + offset, ws.max_column)

    manufacturer = str(field_values.get("제조사", "") or "")
    model_name = str(field_values.get("모델명", "") or "")
    os_version = str(field_values.get("OS버전", "") or "")
    place_name = str(field_values.get("납품장소", "") or "")
    maintenance_end = _parse_excel_date(field_values.get("유지보수종료일"))

    for offset in range(row_count):
        row = base_row + offset
        serial = serial_numbers[offset] if offset < len(serial_numbers) else ""
        ci_cell = ws.cell(row, 2)
        ci_cell.number_format = "@"
        ci_cell.value = str(offset + 1)
        ws.cell(row, 3).value = ""
        ws.cell(row, 4).value = manufacturer
        ws.cell(row, 5).value = model_name
        ws.cell(row, 6).value = serial
        ws.cell(row, 7).value = os_version
        ws.cell(row, 8).value = place_name
        ws.cell(row, 9).value = manufacturer
        ws.cell(row, 10).value = maintenance_end or str(field_values.get("유지보수종료일", "") or "")
        ws.cell(row, 11).value = "Y"

    if "CODELIST" in wb.sheetnames:
        codes = wb["CODELIST"]
        codes["B3"] = model_name
        codes["B168"] = manufacturer

    output = io.BytesIO()
    wb.save(output)
    return output.getvalue()


def _render_delivery_confirmation_xls(template_data: bytes, field_values: dict) -> bytes:
    import xlrd
    from xlutils.copy import copy as xl_copy

    source = xlrd.open_workbook(file_contents=template_data, formatting_info=True)
    source_sheet = source.sheet_by_index(0)
    writable = xl_copy(source)
    ws = writable.get_sheet(0)

    serial_numbers = [serial for serial in _extract_serial_numbers(field_values) if serial]
    xf_text = _get_xls_template_xf_indices(ws, 14, 7)
    xf_footer = _get_xls_template_xf_indices(ws, 31, 7)
    top_date = str(field_values.get("입고일자", "") or "").replace("/", "-")
    footer_date = str(field_values.get("입고일자", "") or "")
    parts = _split_date_parts(field_values.get("입고일자"))
    if parts:
        year, month, day = parts
        footer_date = f"{year}년   {int(month)}월    {int(day)}일"

    values = {
        (6, 2): top_date,
        (10, 2): str(field_values.get("발주명", "") or ""),
        (14, 2): str(field_values.get("모델명", "") or ""),
        (14, 3): "\n".join(serial_numbers),
        (14, 4): str(field_values.get("수량", "") or ""),
        (15, 4): str(field_values.get("수량", "") or ""),
        (16, 4): str(field_values.get("수량", "") or ""),
        (17, 4): str(field_values.get("유지보수수량", "") or field_values.get("수량", "") or ""),
        (31, 4): footer_date,
    }

    for (row_idx, col_idx), value in values.items():
        xf_idx = xf_text[col_idx] if row_idx < 31 else xf_footer[col_idx]
        _set_xls_text_cell(ws, row_idx, col_idx, value, xf_idx)

    _set_xls_multiline_row_height(ws, source_sheet, 14, values[(14, 3)], 14)

    output = io.BytesIO()
    writable.save(output)
    return output.getvalue()


def _extract_purchase_items_for_delivery(field_values: dict) -> list[dict[str, str]]:
    raw_items = field_values.get("__purchase_items")
    normalized: list[dict[str, str]] = []

    if isinstance(raw_items, list):
        for item in raw_items:
            if not isinstance(item, dict):
                continue
            name = str(item.get("name", "") or "").strip()
            quantity = str(item.get("quantity", "") or "").strip()
            unit = str(item.get("unit", "") or "").strip()
            manufacturer = str(item.get("manufacturer", "") or "").strip()
            delivery_place = str(item.get("delivery_place", "") or "").strip()
            if not any([name, quantity, unit, manufacturer, delivery_place]):
                continue
            normalized.append(
                {
                    "name": name,
                    "quantity": quantity,
                    "unit": unit,
                    "manufacturer": manufacturer,
                    "delivery_place": delivery_place,
                }
            )

    if normalized:
        return normalized

    model_name = str(field_values.get("모델명", "") or "").strip()
    quantity = str(field_values.get("수량", "") or "").strip()
    maintenance_quantity = str(field_values.get("유지보수수량", "") or "").strip()
    fallback_items: list[dict[str, str]] = []

    if model_name or quantity:
        fallback_items.append({"name": model_name, "quantity": quantity, "unit": "EA", "manufacturer": "", "delivery_place": ""})
    if maintenance_quantity:
        fallback_items.append(
            {
                "name": f"{model_name} 유지보수".strip(),
                "quantity": maintenance_quantity,
                "unit": "EA",
                "manufacturer": "",
                "delivery_place": "",
            }
        )
    return fallback_items


def _extract_serial_numbers(field_values: dict) -> list[str]:
    serials = field_values.get("__serial_numbers")
    if isinstance(serials, list):
        normalized = [str(serial or "").strip() for serial in serials]
        if normalized:
            return normalized

    fallback = []
    for idx in range(1, 200):
        key = f"시리얼번호{idx}"
        if key not in field_values:
            if idx > 1:
                break
            continue
        fallback.append(str(field_values.get(key, "") or "").strip())

    if fallback:
        return fallback

    return [str(field_values.get("시리얼번호", "") or "").strip()]


def _build_scalar_replacements(field_values: dict, *, serial_separator: str) -> dict[str, object]:
    scalar_values = {
        key: value
        for key, value in field_values.items()
        if isinstance(value, (str, int, float)) or value is None
    }
    serial_numbers = _extract_serial_numbers(field_values)
    filled_serials = [serial for serial in serial_numbers if serial]

    if filled_serials:
        scalar_values["시리얼번호"] = serial_separator.join(filled_serials)
    else:
        scalar_values.setdefault("시리얼번호", "")

    for idx, serial in enumerate(serial_numbers, 1):
        scalar_values[f"시리얼번호{idx}"] = serial
        scalar_values[f"시리얼번호_{idx}"] = serial

    for key in ("발주일자", "입고일자", "검수일자", "유지보수종료일"):
        parts = _split_date_parts(scalar_values.get(key))
        if not parts:
            continue
        year, month, day = parts
        scalar_values[f"{key}_대시"] = f"{year}-{month}-{day}"
        scalar_values[f"{key}_슬래시공백"] = f"{year}/ {month} /{day}"
        scalar_values[f"{key}_한글"] = f"{year}년 {month}월 {day}일"
        scalar_values[f"{key}_한글띄어쓰기"] = f"{year} 년 {month} 월 {day} 일"
        scalar_values[f"{key}_한글여백"] = f"{year}년   {int(month)}월    {int(day)}일"
        scalar_values[f"{key}_US"] = f"{month}/{day}/{year}"

    return {f"{{{{{key}}}}}": value for key, value in scalar_values.items()}


def _build_docx_legacy_date_replacements(field_values: dict) -> dict[str, str]:
    replacements: dict[str, str] = {}

    for key, sample_parts in (
        ("발주일자", ("2026", "03", "19")),
        ("입고일자", ("2026", "03", "30")),
        ("검수일자", ("2026", "03", "31")),
    ):
        parts = _split_date_parts(field_values.get(key))
        if not parts:
            continue
        year, month, day = parts
        sample_year, sample_month, sample_day = sample_parts
        replacements.update(
            {
                f"{sample_year} 년  {sample_month} 월  {sample_day} 일": f"{year} 년  {month} 월  {day} 일",
                f"{sample_year} 년 {sample_month} 월 {sample_day} 일": f"{year} 년 {month} 월 {day} 일",
                f"{sample_year} 년  {sample_month} 월 {sample_day}일": f"{year} 년  {month} 월 {day}일",
                f"{sample_year} 년 {sample_month} 월 {sample_day}일": f"{year} 년 {month} 월 {day}일",
            }
        )

    return replacements


def _replace_docx_paragraph_text(doc, replacements: dict[str, str]) -> None:
    if not replacements:
        return

    for paragraph in _iter_docx_paragraphs(doc):
        full_text = "".join(run.text for run in paragraph.runs)
        if not full_text:
            continue
        updated_text = full_text
        for old, new in replacements.items():
            if old in updated_text:
                updated_text = updated_text.replace(old, new)
        if updated_text == full_text or not paragraph.runs:
            continue
        paragraph.runs[0].text = updated_text
        for run in paragraph.runs[1:]:
            run.text = ""


def _iter_docx_table_paragraphs(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                yield paragraph
            for nested in cell.tables:
                yield from _iter_docx_table_paragraphs(nested)


def _iter_docx_paragraphs(doc):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        yield from _iter_docx_table_paragraphs(table)
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            yield paragraph
        for table in section.header.tables:
            yield from _iter_docx_table_paragraphs(table)
        for paragraph in section.footer.paragraphs:
            yield paragraph
        for table in section.footer.tables:
            yield from _iter_docx_table_paragraphs(table)


def _replace_text_in_docx_element(element, old: str, new: str):
    for node in element.iter():
        if not getattr(node, "tag", "").endswith("}t"):
            continue
        if node.text and old in node.text:
            node.text = node.text.replace(old, new)


def _expand_docx_serial_rows(doc, serial_numbers: list[str]):
    for table in doc.tables:
        row_idx = 0
        while row_idx < len(table.rows):
            row = table.rows[row_idx]
            row_text = " || ".join(cell.text for cell in row.cells)
            if "{{시리얼번호}}" not in row_text:
                row_idx += 1
                continue

            source_tr = row._tr
            parent = source_tr.getparent()
            source_index = parent.index(source_tr)
            template_tr = deepcopy(source_tr)

            for offset, serial in enumerate(serial_numbers):
                if offset == 0:
                    target_tr = source_tr
                else:
                    target_tr = deepcopy(template_tr)
                    parent.insert(source_index + offset, target_tr)
                _replace_text_in_docx_element(target_tr, "{{시리얼번호}}", serial)
                _replace_text_in_docx_element(target_tr, "{{수량}}", "1")

            row_idx += len(serial_numbers)


def _copy_xlsx_row(ws, source_row: int, target_row: int, max_col: int):
    ws.row_dimensions[target_row].height = ws.row_dimensions[source_row].height
    for col in range(1, max_col + 1):
        source = ws.cell(source_row, col)
        target = ws.cell(target_row, col)
        target._style = copy(source._style)
        target.font = copy(source.font)
        target.fill = copy(source.fill)
        target.border = copy(source.border)
        target.alignment = copy(source.alignment)
        target.number_format = source.number_format
        target.protection = copy(source.protection)
        target.value = source.value


def _get_xls_template_xf_indices(ws, template_row_idx: int, max_col: int) -> list[int]:
    row = ws._Worksheet__rows.get(template_row_idx)
    if not row:
        return [17] * max_col

    indices: list[int] = []
    for col_idx in range(max_col):
        cell = row._Row__cells.get(col_idx)
        indices.append(getattr(cell, "xf_idx", row.get_xf_index()))
    return indices


def _set_xls_text_cell(ws, row_idx: int, col_idx: int, value: str, xf_idx: int) -> None:
    from xlwt.Cell import StrCell

    row = ws.row(row_idx)
    sst_idx = ws._Worksheet__parent._Workbook__sst.add_str(value)
    row.insert_cell(col_idx, StrCell(row_idx, col_idx, xf_idx, sst_idx))


def _expand_xlsx_serial_rows(ws, serial_numbers: list[str]):
    row = 1
    max_col = ws.max_column
    while row <= ws.max_row:
        row_values = [ws.cell(row, col).value for col in range(1, max_col + 1)]
        if not any(isinstance(value, str) and "{{시리얼번호}}" in value for value in row_values):
            row += 1
            continue

        if len(serial_numbers) > 1:
            ws.insert_rows(row + 1, len(serial_numbers) - 1)
            for offset in range(1, len(serial_numbers)):
                _copy_xlsx_row(ws, row, row + offset, max_col)

        for offset, serial in enumerate(serial_numbers):
            target_row = row + offset
            for col in range(1, max_col + 1):
                cell = ws.cell(target_row, col)
                if isinstance(cell.value, str):
                    cell.value = cell.value.replace("{{시리얼번호}}", serial)
                    cell.value = cell.value.replace("{{수량}}", "1")

        row += len(serial_numbers)


def _extract_idc_access_people(field_values: dict) -> list[dict[str, str]]:
    people = field_values.get("__idc_access_people")
    if isinstance(people, list):
        normalized = []
        for person in people:
            if not isinstance(person, dict):
                continue
            normalized.append(
                {
                    "company": str(person.get("company", "") or ""),
                    "name": str(person.get("name", "") or ""),
                    "position": str(person.get("position", "") or ""),
                    "contact": str(person.get("contact", "") or ""),
                }
            )
        if normalized:
            return normalized

    fallback = []
    for idx in range(1, 100):
        company = str(field_values.get(f"출입자{idx}_회사명", "") or "")
        name = str(field_values.get(f"출입자{idx}_이름", "") or "")
        position = str(field_values.get(f"출입자{idx}_직책", "") or "")
        contact = str(field_values.get(f"출입자{idx}_연락처", "") or "")
        if not any([company, name, position, contact]):
            if idx > 2:
                break
            continue
        fallback.append(
            {
                "company": company,
                "name": name,
                "position": position,
                "contact": contact,
            }
        )

    return fallback or [{"company": "", "name": "", "position": "", "contact": ""}]


def _copy_row_style(ws, source_row: int, target_row: int, max_col: int = 5):
    ws.row_dimensions[target_row].height = ws.row_dimensions[source_row].height
    for col in range(1, max_col + 1):
        source = ws.cell(source_row, col)
        target = ws.cell(target_row, col)
        target._style = copy(source._style)
        target.font = copy(source.font)
        target.fill = copy(source.fill)
        target.border = copy(source.border)
        target.alignment = copy(source.alignment)
        target.number_format = source.number_format
        target.protection = copy(source.protection)


def _multiline_row_height(base_height: float | int | None, text: str) -> float:
    base = float(base_height) if base_height else 15.0
    line_count = max(1, len(str(text or "").splitlines()))
    return base if line_count <= 1 else max(base, base * line_count)


def _set_xls_multiline_row_height(ws, source_sheet, row_idx: int, text: str, reference_row_idx: int) -> None:
    row = ws.row(row_idx)
    source_info = source_sheet.rowinfo_map.get(reference_row_idx)
    base_height = source_info.height if source_info and source_info.height else 255
    row.height_mismatch = True
    row.height = int(_multiline_row_height(base_height, text))


def _render_idc_access_xlsx(template_data: bytes, field_values: dict) -> bytes:
    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(template_data))
    ws = wb.active
    people = _extract_idc_access_people(field_values)
    start_row = 10
    end_row = max(40, start_row + len(people) - 1)

    for row in range(start_row, end_row + 1):
        _copy_row_style(ws, 10, row)
        for col in range(1, 6):
            ws.cell(row, col).value = ""

    for idx, person in enumerate(people):
        row = start_row + idx
        ws.cell(row, 1).value = ""
        ws.cell(row, 2).value = person["company"]
        ws.cell(row, 3).value = person["name"]
        ws.cell(row, 4).value = person["position"]
        ws.cell(row, 5).value = person["contact"]

    output = io.BytesIO()
    wb.save(output)
    return output.getvalue()


def _find_delivery_confirmation_footer_row(ws) -> int:
    for row in range(1, ws.max_row + 1):
        for col in range(1, min(ws.max_column, 8) + 1):
            value = ws.cell(row, col).value
            if isinstance(value, str) and "납품 / 인수 일" in value:
                return row
    return 32


def _render_delivery_confirmation_xlsx(template_data: bytes, field_values: dict) -> bytes:
    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(template_data))
    ws = wb[wb.sheetnames[0]]

    purchase_items = _extract_purchase_items_for_delivery(field_values)
    serial_numbers = [serial for serial in _extract_serial_numbers(field_values) if serial]
    top_date = str(field_values.get("입고일자", "") or "").replace("/", "-")
    footer_date = str(field_values.get("입고일자_한글여백", "") or "")
    if not footer_date:
        parts = _split_date_parts(field_values.get("입고일자"))
        if parts:
            year, month, day = parts
            footer_date = f"{year}년   {int(month)}월    {int(day)}일"

    item_start_row = 15
    footer_row = _find_delivery_confirmation_footer_row(ws)
    template_row = max(item_start_row, footer_row - 1)
    available_rows = max(footer_row - item_start_row, 1)
    row_count = max(len(purchase_items), available_rows)

    if len(purchase_items) > available_rows:
        extra_rows = len(purchase_items) - available_rows
        ws.insert_rows(footer_row, extra_rows)
        for offset in range(extra_rows):
            _copy_xlsx_row(ws, template_row, footer_row + offset, ws.max_column)
        footer_row += extra_rows
        row_count = len(purchase_items)

    ws["C7"] = top_date
    ws["C11"] = str(field_values.get("발주명", "") or "")

    for row in range(item_start_row, item_start_row + row_count):
        ws.cell(row, 2).value = row - item_start_row + 1
        ws.cell(row, 3).value = ""
        ws.cell(row, 4).value = ""
        ws.cell(row, 5).value = ""

    for index, item in enumerate(purchase_items):
        row = item_start_row + index
        ws.cell(row, 2).value = index + 1
        ws.cell(row, 3).value = item["name"]
        ws.cell(row, 4).value = "\n".join(serial_numbers) if index == 0 else ""
        ws.cell(row, 5).value = item["quantity"]

    serial_text = "\n".join(serial_numbers)
    base_height = ws.row_dimensions[item_start_row].height or ws.row_dimensions[template_row].height
    ws.row_dimensions[item_start_row].height = _multiline_row_height(base_height, serial_text)

    ws.cell(footer_row, 5).value = footer_date

    output = io.BytesIO()
    wb.save(output)
    return output.getvalue()


def _render_idc_access_xls(template_data: bytes, field_values: dict) -> bytes:
    import xlrd
    from xlutils.copy import copy as xl_copy

    source = xlrd.open_workbook(file_contents=template_data, formatting_info=True)
    writable = xl_copy(source)
    ws = writable.get_sheet(0)
    people = _extract_idc_access_people(field_values)
    start_row = 9
    clear_until = max(start_row + len(people) - 1, 39)
    xf_indices = _get_xls_template_xf_indices(ws, start_row, 5)

    ws.set_panes_frozen(True)
    ws.set_horz_split_pos(8)
    ws.set_horz_split_first_visible(8)
    ws.set_remove_splits(True)

    for row_idx in range(start_row, clear_until + 1):
        person_idx = row_idx - start_row
        person = people[person_idx] if person_idx < len(people) else None
        values = (
            [
                "",
                person["company"],
                person["name"],
                person["position"],
                person["contact"],
            ]
            if person
            else ["", "", "", "", ""]
        )
        for col_idx, value in enumerate(values):
            _set_xls_text_cell(ws, row_idx, col_idx, value, xf_indices[col_idx])

    output = io.BytesIO()
    writable.save(output)
    return output.getvalue()


def _is_dynamic_idc_key(key: str) -> bool:
    return bool(re.fullmatch(r"출입자\d+_(회사명|이름|직책|연락처)", key))


@router.post("/{bundle_id}/purchase-order/extract", response_model=BundlePurchaseOrderExtractResponse)
async def extract_purchase_order(
    bundle_id: int,
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    _: User = Depends(get_current_user),
):
    bundle = db.query(TemplateBundle).filter(TemplateBundle.id == bundle_id).first()
    if not bundle:
        raise HTTPException(status_code=404, detail="번들을 찾을 수 없습니다")

    filename = file.filename or "purchase-order.pdf"
    if not filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="PDF 파일만 업로드할 수 있습니다")

    data = await file.read()
    if not data:
        raise HTTPException(status_code=400, detail="업로드된 파일이 비어 있습니다")

    try:
        parsed = parse_purchase_order_pdf(data)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except RuntimeError as e:
        raise HTTPException(status_code=500, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"발주서 파싱 오류: {e}")

    bundle_keys = {
        variable.get("key")
        for variable in (bundle.variables or [])
        if isinstance(variable, dict) and variable.get("key")
    }
    field_values = {
        key: value
        for key, value in parsed["field_values"].items()
        if not bundle_keys or key in bundle_keys
    }

    extracted_keys = [key for key in parsed["extracted_keys"] if not bundle_keys or key in bundle_keys]
    inferred_keys = [key for key in parsed["inferred_keys"] if not bundle_keys or key in bundle_keys]
    missing_keys = (
        sorted(key for key in (bundle_keys - set(field_values.keys())) if not _is_dynamic_idc_key(key))
        if bundle_keys
        else []
    )

    return BundlePurchaseOrderExtractResponse(
        filename=filename,
        field_values=field_values,
        purchase_items=parsed.get("purchase_items", []),
        extracted_keys=extracted_keys,
        inferred_keys=inferred_keys,
        missing_keys=missing_keys,
        warnings=parsed["warnings"],
    )


@router.post("/{bundle_id}/generate")
async def generate_bundle(
    bundle_id: int,
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    b = db.query(TemplateBundle).filter(TemplateBundle.id == bundle_id).first()
    if not b:
        raise HTTPException(status_code=404, detail="번들을 찾을 수 없습니다")
    b = _sync_bundle_items_from_folder(db, b)

    form = await request.form()
    try:
        field_values = json.loads(str(form.get("field_values", "{}")))
        selected_items = json.loads(str(form.get("selected_items", "[]")))
    except (ValueError, json.JSONDecodeError) as e:
        raise HTTPException(status_code=422, detail=f"잘못된 요청: {e}")

    if not selected_items:
        selected_items = [item.id for item in _ordered_bundle_items(b)]

    items = [item for item in _ordered_bundle_items(b) if item.id in selected_items]
    if not items:
        raise HTTPException(status_code=400, detail="생성할 문서가 선택되지 않았습니다")

    zip_buffer = io.BytesIO()
    records_created = []

    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        for item in items:
            tpl = item.template
            if not tpl:
                continue

            template_data = minio_client.download_file(tpl.minio_object_key)
            if template_data is None:
                continue

            try:
                if item.display_name == "IDC 출입명단" and tpl.file_type == "xls":
                    rendered = _render_idc_access_xls(template_data, field_values)
                elif item.display_name == "IDC 출입명단" and tpl.file_type == "xlsx":
                    rendered = _render_idc_access_xlsx(template_data, field_values)
                elif item.display_name == "납품확인서" and tpl.file_type == "xls":
                    rendered = _render_delivery_confirmation_xls(template_data, field_values)
                elif item.display_name == "납품확인서" and tpl.file_type == "xlsx":
                    rendered = _render_delivery_confirmation_xlsx(template_data, field_values)
                elif item.display_name == "CMDB" and tpl.file_type == "xlsm":
                    rendered = _render_cmdb_xlsm(template_data, field_values)
                elif tpl.file_type == "docx":
                    rendered = _render_docx(template_data, field_values)
                else:
                    rendered = _render_xlsx(
                        template_data,
                        field_values,
                        expand_serial_rows=item.display_name != "납품확인서",
                        keep_vba=tpl.file_type == "xlsm",
                    )
            except Exception as e:
                print(f"[Bundle] 렌더링 오류 ({item.display_name}): {e}")
                continue

            # Build output filename from pattern
            output_name = _resolve_output_name(item, field_values)
            safe_name = _sanitize_filename(output_name)
            filename = f"{safe_name}.{tpl.file_type}"

            zf.writestr(filename, rendered)

            # Save as record
            object_key = f"documents/records/{current_user.id}/{filename}"
            minio_client.upload_file(object_key, rendered, _content_type(tpl.file_type))

            record = DocumentRecord(
                template_id=tpl.id,
                title=f"{output_name}",
                field_values=field_values,
                original_filename=filename,
                minio_object_key=object_key,
                file_size=len(rendered),
                created_by=current_user.id,
            )
            db.add(record)
            records_created.append(record)

    db.commit()

    zip_data = zip_buffer.getvalue()
    bundle_title = field_values.get("발주명", b.name)
    safe_title = _sanitize_filename(bundle_title)
    zip_filename = f"{safe_title}_문서일괄.zip"

    return Response(
        content=zip_data,
        media_type="application/zip",
        headers={
            "Content-Disposition": f"attachment; filename*=UTF-8''{quote(zip_filename)}",
        },
    )
