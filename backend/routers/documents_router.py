import io
import json
import re
import zipfile
from typing import List
from urllib.parse import quote

from fastapi import APIRouter, Depends, File, Form, HTTPException, Request, UploadFile
from fastapi.responses import Response
from sqlalchemy.orm import Session
from sqlalchemy.exc import IntegrityError

import minio_client
from auth import get_current_user
from database import get_db
from models import DocumentRecord, DocumentTemplate, TemplateBundleItem, User
from schemas import (
    DocumentRecordResponse,
    DocumentTemplateResponse,
    DocumentTemplateUpdate,
)

router = APIRouter(prefix="/api/documents", tags=["documents"])

_PLACEHOLDER_RE = re.compile(r"\{\{([^{}]+)\}\}")

DOCX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
XLSX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
)
XLSM_CONTENT_TYPE = "application/vnd.ms-excel.sheet.macroEnabled.12"
XLS_CONTENT_TYPE = "application/vnd.ms-excel"


def _content_type(file_type: str) -> str:
    return {
        "docx": DOCX_CONTENT_TYPE,
        "xlsx": XLSX_CONTENT_TYPE,
        "xlsm": XLSM_CONTENT_TYPE,
        "xls": XLS_CONTENT_TYPE,
    }.get(file_type, "application/octet-stream")


# ── 변수 추출 ────────────────────────────────────────────────────────────────

def _extract_variables_docx(data: bytes) -> list[dict]:
    from docx import Document

    doc = Document(io.BytesIO(data))
    variables: dict[str, dict] = {}

    def scan(text: str):
        for m in _PLACEHOLDER_RE.finditer(text):
            key = m.group(1).strip()
            if key not in variables:
                variables[key] = {"key": key, "label": key, "type": "text",
                                  "img_width": None, "img_height": None}

    for para in doc.paragraphs:
        scan(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    scan(para.text)

    return sorted(variables.values(), key=lambda v: v["key"])


def _extract_variables_xlsx(data: bytes) -> list[dict]:
    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(data))
    variables: dict[str, dict] = {}

    for sheet in wb.worksheets:
        for row in sheet.iter_rows():
            for cell in row:
                if cell.value and isinstance(cell.value, str):
                    for m in _PLACEHOLDER_RE.finditer(cell.value):
                        key = m.group(1).strip()
                        if key not in variables:
                            variables[key] = {"key": key, "label": key, "type": "text",
                                              "img_width": None, "img_height": None}

    return sorted(variables.values(), key=lambda v: v["key"])


# ── 렌더링 ──────────────────────────────────────────────────────────────────

def _render_docx(
    template_data: bytes,
    field_values: dict,
    image_files: dict[str, bytes] | None = None,
    image_vars: list[dict] | None = None,
) -> bytes:
    from docxtpl import DocxTemplate, InlineImage
    from docx.shared import Mm

    tpl = DocxTemplate(io.BytesIO(template_data))
    context = dict(field_values)

    if image_files and image_vars:
        def to_mm(val: float, unit: str) -> float:
            return val * 10 if unit == "cm" else val

        img_meta = {
            v["key"]: (
                to_mm(v["img_width"], v.get("img_unit") or "mm"),
                to_mm(v["img_height"], v.get("img_unit") or "mm"),
            )
            for v in image_vars
            if v.get("type") == "image"
            and v.get("img_width") and v.get("img_height")
        }
        for key, (width_mm, height_mm) in img_meta.items():
            img_data = image_files.get(key)
            if not img_data:
                continue
            try:
                from PIL import Image as PILImage
                pil_img = PILImage.open(io.BytesIO(img_data))
                px_w = round(width_mm / 25.4 * 96)
                px_h = round(height_mm / 25.4 * 96)
                pil_img = pil_img.resize((px_w, px_h), PILImage.LANCZOS)
                # convert first, then choose format based on final mode
                if pil_img.mode not in ("RGB", "RGBA"):
                    pil_img = pil_img.convert("RGB")
                fmt = "PNG" if pil_img.mode == "RGBA" else "JPEG"
                img_buf = io.BytesIO()
                pil_img.save(img_buf, format=fmt)
                img_buf.seek(0)
            except ImportError:
                # Pillow 미설치 시 원본 이미지를 그대로 사용
                img_buf = io.BytesIO(img_data)
                img_buf.seek(0)
            except Exception:
                img_buf = io.BytesIO(img_data)
                img_buf.seek(0)
            context[key] = InlineImage(tpl, img_buf, width=Mm(width_mm), height=Mm(height_mm))

    tpl.render(context)
    output = io.BytesIO()
    tpl.save(output)
    return output.getvalue()


def _render_xlsx(template_data: bytes, field_values: dict) -> bytes:
    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(template_data))

    for sheet in wb.worksheets:
        for row in sheet.iter_rows():
            for cell in row:
                if cell.value and isinstance(cell.value, str):
                    val = cell.value
                    for key, value in field_values.items():
                        val = val.replace(f"{{{{{key}}}}}", str(value))
                    cell.value = val

    output = io.BytesIO()
    wb.save(output)
    return output.getvalue()


# ── 헬퍼 ─────────────────────────────────────────────────────────────────────

def _template_to_response(t: DocumentTemplate) -> DocumentTemplateResponse:
    return DocumentTemplateResponse(
        id=t.id,
        name=t.name,
        description=t.description,
        file_type=t.file_type,
        original_filename=t.original_filename,
        file_size=t.file_size,
        variables=t.variables or [],
        created_at=t.created_at,
        created_by_username=t.creator.username if t.creator else None,
    )


# ── 템플릿 엔드포인트 ─────────────────────────────────────────────────────────

@router.get("/templates", response_model=List[DocumentTemplateResponse])
def list_templates(
    db: Session = Depends(get_db),
    _: User = Depends(get_current_user),
):
    templates = (
        db.query(DocumentTemplate)
        .order_by(DocumentTemplate.created_at.desc())
        .all()
    )
    return [_template_to_response(t) for t in templates]


@router.post("/templates", response_model=DocumentTemplateResponse)
async def create_template(
    file: UploadFile = File(...),
    name: str = Form(...),
    description: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    filename = file.filename or "template"
    ext = filename.rsplit(".", 1)[-1].lower()
    if ext not in ("docx", "xlsx"):
        raise HTTPException(status_code=400, detail="docx 또는 xlsx 파일만 지원합니다")

    data = await file.read()

    try:
        variables = (
            _extract_variables_docx(data)
            if ext == "docx"
            else _extract_variables_xlsx(data)
        )
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"파일 파싱 오류: {e}")

    object_key = f"documents/templates/{current_user.id}/{filename}"
    if not minio_client.upload_file(object_key, data, _content_type(ext)):
        raise HTTPException(status_code=500, detail="파일 업로드 실패")

    template = DocumentTemplate(
        name=name,
        description=description or None,
        file_type=ext,
        original_filename=filename,
        minio_object_key=object_key,
        file_size=len(data),
        variables=variables,
        created_by=current_user.id,
    )
    db.add(template)
    db.commit()
    db.refresh(template)
    return _template_to_response(template)


@router.get("/templates/{template_id}", response_model=DocumentTemplateResponse)
def get_template(
    template_id: int,
    db: Session = Depends(get_db),
    _: User = Depends(get_current_user),
):
    t = db.query(DocumentTemplate).filter(DocumentTemplate.id == template_id).first()
    if not t:
        raise HTTPException(status_code=404, detail="템플릿을 찾을 수 없습니다")
    return _template_to_response(t)


@router.put("/templates/{template_id}", response_model=DocumentTemplateResponse)
def update_template(
    template_id: int,
    body: DocumentTemplateUpdate,
    db: Session = Depends(get_db),
    _: User = Depends(get_current_user),
):
    t = db.query(DocumentTemplate).filter(DocumentTemplate.id == template_id).first()
    if not t:
        raise HTTPException(status_code=404, detail="템플릿을 찾을 수 없습니다")
    if body.name is not None:
        t.name = body.name
    if body.description is not None:
        t.description = body.description
    if body.variables is not None:
        t.variables = [
            {
                "key": v.key,
                "label": v.label,
                "type": v.type,
                "img_width": v.img_width,
                "img_height": v.img_height,
                "img_unit": v.img_unit or "mm",
            }
            for v in body.variables
        ]
    db.commit()
    db.refresh(t)
    return _template_to_response(t)


@router.delete("/templates/{template_id}", status_code=204)
def delete_template(
    template_id: int,
    db: Session = Depends(get_db),
    _: User = Depends(get_current_user),
):
    t = db.query(DocumentTemplate).filter(DocumentTemplate.id == template_id).first()
    if not t:
        raise HTTPException(status_code=404, detail="템플릿을 찾을 수 없습니다")

    bundle_items = (
        db.query(TemplateBundleItem)
        .filter(TemplateBundleItem.template_id == template_id)
        .all()
    )
    if bundle_items:
        bundle_names = ", ".join(item.display_name for item in bundle_items)
        raise HTTPException(
            status_code=409,
            detail=(
                "이 템플릿은 번들에서 사용 중이라 삭제할 수 없습니다. "
                f"연결된 번들 항목: {bundle_names}"
            ),
        )

    object_key = t.minio_object_key
    db.delete(t)
    try:
        db.commit()
    except IntegrityError:
        db.rollback()
        raise HTTPException(
            status_code=409,
            detail="다른 문서 또는 번들에서 사용 중인 템플릿이라 삭제할 수 없습니다.",
        )

    minio_client.delete_file(object_key)


@router.get("/templates/{template_id}/file")
def download_template_file(
    template_id: int,
    db: Session = Depends(get_db),
    _: User = Depends(get_current_user),
):
    t = db.query(DocumentTemplate).filter(DocumentTemplate.id == template_id).first()
    if not t:
        raise HTTPException(status_code=404, detail="템플릿을 찾을 수 없습니다")
    data = minio_client.download_file(t.minio_object_key)
    if data is None:
        raise HTTPException(status_code=404, detail="파일을 찾을 수 없습니다")
    encoded_name = quote(t.original_filename)
    return Response(
        content=data,
        media_type=_content_type(t.file_type),
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_name}"},
    )


# ── 작업 기록 엔드포인트 ──────────────────────────────────────────────────────

@router.get("/records", response_model=List[DocumentRecordResponse])
def list_records(
    db: Session = Depends(get_db),
    _: User = Depends(get_current_user),
):
    records = (
        db.query(DocumentRecord)
        .order_by(DocumentRecord.created_at.desc())
        .all()
    )
    return [
        DocumentRecordResponse(
            id=r.id,
            template_id=r.template_id,
            template_name=r.template.name if r.template else None,
            file_type=r.template.file_type if r.template else None,
            title=r.title,
            field_values=r.field_values or {},
            original_filename=r.original_filename,
            file_size=r.file_size,
            created_at=r.created_at,
            created_by_username=r.creator.username if r.creator else None,
        )
        for r in records
    ]


@router.post("/records", response_model=DocumentRecordResponse)
async def create_record(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    form = await request.form()

    try:
        template_id = int(form["template_id"])
        title = str(form["title"])
        field_values = json.loads(str(form.get("field_values", "{}")))
    except (KeyError, ValueError) as e:
        raise HTTPException(status_code=422, detail=f"잘못된 요청: {e}")

    t = db.query(DocumentTemplate).filter(DocumentTemplate.id == template_id).first()
    if not t:
        raise HTTPException(status_code=404, detail="템플릿을 찾을 수 없습니다")

    known_image_keys = {
        v["key"] for v in (t.variables or []) if v.get("type") == "image"
    }
    image_files: dict[str, bytes] = {}
    for key, value in form.multi_items():
        if key.startswith("img__") and hasattr(value, "read"):
            var_key = key[5:]
            if var_key in known_image_keys:
                image_files[var_key] = await value.read()

    template_data = minio_client.download_file(t.minio_object_key)
    if template_data is None:
        raise HTTPException(status_code=404, detail="템플릿 파일을 찾을 수 없습니다")

    try:
        rendered = (
            _render_docx(template_data, field_values, image_files, t.variables or [])
            if t.file_type == "docx"
            else _render_xlsx(template_data, field_values)
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"문서 렌더링 오류: {e}")

    safe_title = re.sub(r'[\\/*?:"<>|]', "_", title)
    output_filename = f"{safe_title}.{t.file_type}"
    object_key = f"documents/records/{current_user.id}/{output_filename}"

    if not minio_client.upload_file(object_key, rendered, _content_type(t.file_type)):
        raise HTTPException(status_code=500, detail="파일 저장 실패")

    record = DocumentRecord(
        template_id=t.id,
        title=title,
        field_values=field_values,
        original_filename=output_filename,
        minio_object_key=object_key,
        file_size=len(rendered),
        created_by=current_user.id,
    )
    db.add(record)
    db.commit()
    db.refresh(record)

    return DocumentRecordResponse(
        id=record.id,
        template_id=record.template_id,
        template_name=t.name,
        file_type=t.file_type,
        title=record.title,
        field_values=record.field_values or {},
        original_filename=record.original_filename,
        file_size=record.file_size,
        created_at=record.created_at,
        created_by_username=current_user.username,
    )


@router.delete("/records/{record_id}", status_code=204)
def delete_record(
    record_id: int,
    db: Session = Depends(get_db),
    _: User = Depends(get_current_user),
):
    r = db.query(DocumentRecord).filter(DocumentRecord.id == record_id).first()
    if not r:
        raise HTTPException(status_code=404, detail="작업 기록을 찾을 수 없습니다")
    minio_client.delete_file(r.minio_object_key)
    db.delete(r)
    db.commit()


@router.get("/records/{record_id}/file")
def download_record_file(
    record_id: int,
    db: Session = Depends(get_db),
    _: User = Depends(get_current_user),
):
    r = db.query(DocumentRecord).filter(DocumentRecord.id == record_id).first()
    if not r:
        raise HTTPException(status_code=404, detail="작업 기록을 찾을 수 없습니다")
    data = minio_client.download_file(r.minio_object_key)
    if data is None:
        raise HTTPException(status_code=404, detail="파일을 찾을 수 없습니다")
    file_type = r.template.file_type if r.template else (r.original_filename or "").rsplit(".", 1)[-1]
    encoded_name = quote(r.original_filename or "document")
    return Response(
        content=data,
        media_type=_content_type(file_type),
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_name}"},
    )
